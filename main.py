import PyPDF2
import os, shutil
import openai
from openai import Embedding
import docx
from langdetect import detect
import time, re
import requests
import streamlit as st
from docx import Document
import base64, pyperclip
from langdetect.lang_detect_exception import LangDetectException

api_key = '******************************'
openai.api_key = api_key


books = []
tran = []
pages = []
pages1 = []

#books code
def textdatab():
  base_path = 'books'               
  author = 'Maulana Wahiduddin Khan'                          
  for name in os.listdir(base_path):                     
    print(name)
    book_name = name.split('.')[0].split('_')[1]                      
    book = { 'Author':author, 'Book Name':book_name, 'Content':[] }                  
    file_path = os.path.join(base_path,name)                               
    pdfFileObj = open(file_path, 'rb')                    
    pdfReader = PyPDF2.PdfReader(pdfFileObj)                         

    add = False                               
    temp = {'page':'', 'chapter':'', 'content':''}                     
    for num_p,page in enumerate(pdfReader.pages[1:]):                                   

      text = page.extract_text()                                
      split_text = text.split('\n')                
      text_length = len(split_text)                
      text_length_flag = text_length>3                           
      page_length = len(pdfReader.pages)                             
      if text_length_flag:                      
        chap = " ".join(str(split_text[1]).lower().strip().split())                 
        chap_2 = " ".join(str(split_text[2]).lower().strip().split())                 
        bk = " ".join(str(split_text[0]).lower().strip().split())                 

      if (text_length_flag):                     
        if ( (chap!='introduction') and (chap!='table of contents') and (chap!='foreword') and (chap!='table of content') and (chap!='preface') and (chap!='forward') and (chap!='publisher’s note') ):                    
          add = True                  
        else:               
          continue                 
      if (text_length_flag):             
        if ( ((num_p>(page_length*(3/4)))and(bk==chap)) or ((num_p>(page_length*(3/4)))and(chap in bk)and(chap_2 in bk)) or ((num_p>(page_length*(3/4)))and(chap=='a final word')) or (('conclusion'in chap)and(num_p>(page_length*(3/4)))) or (chap=='notes') or (chap=='index') or (chap=='in search of god') or (chap=='the callof the quran') ):                          
          break                    
      if ( add and (text_length_flag) ):                                  
        temp['page'] = str(int(num_p+2))                   
        temp['chapter'] = chap                         
                                      
        for num,txt in enumerate(split_text[3:]):                  
          a = " ".join(str(txt).strip().split())                     
          a1 = ''                                             
          for char in txt:                         
            if char.isalpha():                 
              a1 += char                           
          if a!='' and a!=' ' and not(a1.isupper()):                     
            if not( (num==1 or num==2) and ( (" ".join(str(txt).lower().strip().split())) in temp['chapter'] ) ):  
              temp['content'] += a+' '                                                              
                                                
        if len(temp['content']):                               
          book['Content'].append(temp)                        
        temp = {'page':'', 'chapter':'', 'content':''}     
        break                  
    
    books.append(book)

def read_docx(file_path):
    pages = []  # Initialize an empty list to store paragraphs
    doc = docx.Document(file_path)
    
    paragraphs = doc.paragraphs
    
    for i in range(len(paragraphs)):
        current_paragraph = paragraphs[i].text  # Access the text property of the paragraph
        print(current_paragraph)
        
        # Your original condition to exclude paragraphs ending with numbers
        if (current_paragraph and not current_paragraph.endswith(tuple(str(i) for i in range(9999)))):
            # Additional condition to exclude "Printed in India" and limit to 50 paragraphs
            if current_paragraph.strip() != "Printed in India" or i >= 50:
                pages.append(current_paragraph)

    print(pages)
    return pages

# def chatgpt_query(question, contexts):
#     messages = [{"role": "system", "content": contexts}, {"role": "user", "content": question}]
#     response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=messages)
#     reply = response.choices[0].text
#     return reply
    
def chatgpt_query(context, writer, prompt):
    print('hello jee')
    if not context.strip():
        return context

    try:
        lang = detect(context)
        if lang in ['en', 'hi']:
            return context
    except LangDetectException as e:
        print("Language detection error:", str(e))
        return 0

    try:
        # instruction_message = f"Assistant is an intelligent chatbot which translates Urdu into English.\nInstructions:\n- Translate it while preserving the tone of {writer}"
        messages = [
            {"role": "system", "content": prompt},
            {"role": "user", "content": context}
        ]
        response = openai.ChatCompletion.create(model="gpt-3.5-turbo-16k", messages=messages)
        print("Translated")
        return response.choices[0].message.content
    except openai.error.InvalidRequestError as e:
        print("Length exceeds")
        return 0


def get_info(context,content):
    print('hello jee')
    if not context.strip():
      return context
    
    try:
        lang = detect(context)
    except LangDetectException as e:
        print("Language detection error:", str(e))
        return 0
    
    lang = detect(context)
    if lang in ['en', 'hindi']:
        return context
    try:
      response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=f"translate the data into English \ndata:{context}\nAnswer:",
        max_tokens=1000,
        temperature=0,
      )
    except openai.error.InvalidRequestError as e:
       print("Length Exceeds")
       return 0
    print(response.choices[0].text)
    return response.choices[0].text

# def create_book():
#   if "companies" in os.listdir("."):shutil.rmtree("companies")
#   os.mkdir("companies")
#   with open('companies/test1.txt', 'w', encoding='utf-8') as f:
#       for paragraph in tran:
#           f.write(paragraph + '\n')
#   tran.clear()



def create_book():
    if "companies" in os.listdir("."):
        shutil.rmtree("companies")
    os.mkdir("companies")

    doc = Document()
    for paragraph in tran:
        doc.add_paragraph(paragraph)
    tran.clear()
    file_path = 'companies/Translation.docx'
    doc.save(file_path)



def extract(name, prompt):
  # return 0
  # Usage example
  # file_path = f"companies/{name}"
  file_path = name
  pages = read_docx(file_path)
  

  textdatab()
  for book in books:
      for content in book["Content"]:
          Content = content["content"]
          # print(Content)
          break
      
  sentence_chunk = ''
  n = 8  
  for i in range(len(pages)):
    if (i+1)%n==0: 
        x = sentence_chunk + pages[i]
        if len(x) <= 8600:
           sentence_chunk += pages[i]
           pages1.append(sentence_chunk)
           sentence_chunk = ''
        else:
          pages1.append(sentence_chunk)
          sentence_chunk =  pages[i]
    else:
        x = sentence_chunk + pages[i] + '\n'
        if len(x) <= 8600:
           sentence_chunk = sentence_chunk + pages[i] + '\n'
        else:
           pages1.append(sentence_chunk)
           sentence_chunk =  pages[i]
           
  if sentence_chunk:
    pages1.append(sentence_chunk)

  # print(len(pages))
  # print(len(pages1))

  for i, page in enumerate(pages1):
    # if i == 1:
    #     break
    # print(page)
    # x = get_info(page, Content)
    x = chatgpt_query(page, Content, prompt)
    if x != 0:
      tran.append(x)
    # tran.append('\n')
  pages1.clear()
  create_book()
  return "Done"
  # count = ""
  # for i, page in enumerate(pages):
      
  #     # if i <= 32:
  #     #   continue
  #     # print(i)
  #     print(f"Page {i+1}:\n{page}\n")
  #     Content = page

  #     if i == 10:
  #       break
  #     if i == 0:
  #       count = page
  #       continue
  #     if i%2 != 1:
  #        count = ""
  #        count = page

  #     elif i%2 == 0:
  #       if count:
  #          count = count + "\n" + page
  #       else:
  #          count = page

  #       # x = get_info(page, Content)
  #       print(count)
  #       tran.append(count)
  #       tran.append('\n')
  
  create_book()
  return "done"
      

# Function to download a file
def download_file(file_path, file_name):
    with open(file_path, 'rb') as f:
        file_content = f.read()
    st.download_button(label='Download File', data=file_content, file_name=file_name)


def get_download_link(file_path):
    with open(file_path, 'rb') as f:
        file_content = f.read()
    file_name = os.path.basename(file_path)
    encoded_content = base64.b64encode(file_content).decode('utf-8')
    href = f'<a href="data:application/octet-stream;base64,{encoded_content}" download="{file_name}">Click here to download the file</a>'
    return href

def get_improved_file_download_link(file_path):
    with open(file_path, 'rb') as f:
        file_content = f.read()
    
    file_name = os.path.basename(file_path)
    encoded_content = base64.b64encode(file_content).decode('utf-8')
    href = f'<a href="data:application/octet-stream;base64,{encoded_content}" download="{file_name}">Download Improved File</a>'
    return href

def copy_prompt(default_prompt):
    copy_link = f'<a href="#" onclick="navigator.clipboard.writeText(\'{default_prompt}\'); return false;">Copy Default Prompt</a>'
    return copy_link

def improve_english(file_path):
  # Read the content of the selected file
  with open(file_path, 'r', encoding='utf-8') as f:
    content = f.read()

  # Define a prompt for improving English
  improve_english_prompt = "Assistant is here to help you improve your English. Please provide an English translation that you think is better than the current one."

  # Streamlit input to allow the user to provide an improved translation
  improved_translation = st.text_area("Provide an improved English translation:", value=content)

  # If the user provides an improved translation, update the content
  if improved_translation:
    content = improved_translation

  # Save the improved content to a new file
  improved_file_path = file_path.replace('.docx', '_improved.docx')
  with open(improved_file_path, 'w', encoding='utf-8') as f:
    f.write(content)

  return improved_file_path

# Streamlit app code
def main():
    
    # Streamlit code for language selection and file upload
    # st.header('Language Preference')
    # option = st.selectbox('Which language do you prefer?', ["English", "Urdu"])
    # st.write('You selected:', option)

    st.header('Header Prompt')
    default_prompt = "Assistant is an intelligent chatbot which translates Urdu into English."
    prompt = st.text_input("You can change the prompt here", default_prompt)
    option = st.selectbox('You can choose the default prompt from markdown', [prompt, default_prompt])
    # print(option)
    if not prompt:
        option = default_prompt


    
    st.header('Single File Upload')
    uploaded_file = st.file_uploader('Upload a file')

    if uploaded_file is not None:
        if not uploaded_file.name.endswith('.docx'):
            st.write("Enter a docx File")
        else:
            # Save the uploaded file
            if "companies" in os.listdir("."):
                shutil.rmtree("companies")
            os.mkdir("companies")
            file_path = os.path.join('companies', uploaded_file.name)
            # file_path = 'C:\\Users\\hp\\Desktop\\Islam-Stream\\companies\\Rtest.docx'

            print(file_path)
            with open(file_path, 'wb') as f:
                f.write(uploaded_file.read())

            # Call the extract function to process the file
            result = extract(file_path, option)

            # Display the result
            st.header('Download File')
            # Get the list of files in the "companies" folder
            file_list = os.listdir('companies')
            # Display the file list
            selected_file = st.selectbox('Select a file to download:', file_list)
            # Check if a file is selected
            if selected_file:
                file_path = os.path.join('companies', selected_file)
                # download_file(file_path, selected_file)
                st.markdown(get_download_link(file_path), unsafe_allow_html=True)

            st.write("Translated")
            # After the translation is done, provide a button to improve English
            if st.button("Improve English of Translated File"):
              # Get the list of files in the "companies" folder
              file_list = os.listdir('companies')
              
              for selected_file in file_list:
                  file_path = os.path.join('companies', selected_file)
                  
                  # Call a function to improve English using the selected file and a prompt
                  improved_file_path = read_docx(file_path)
                  prompt = "Assistant is here to help you improve your English. Please provide an English translation that you think is better than the current one."
                  option = prompt
                  # Call the extract function to process the file (if needed)
                  result = extract(file_path, option)
                  print(file_path)
                  # Create a download link for the improved file and display it
                  improved_link = get_improved_file_download_link(file_path)
                  st.markdown(improved_link, unsafe_allow_html=True)
                  
                  st.write("English improved successfully!")

                  # st.write("Translated")


                  # st.write("English improved successfully!")

                  # # Display a button to download the improved file
                  # st.header('Download Improved File')
                  # st.markdown(get_download_link(improved_file_path), unsafe_allow_html=True)


if __name__ == '__main__':
    main()


